import os
import random
from flask import Flask, render_template, url_for, redirect, url_for, request, flash, session, jsonify, send_file
from werkzeug.utils import secure_filename
from summary import summarize, getPercentage, getTopWords
from webScrape import getArticle
from readfile import getText
from docx import Document

files = []

UPLOAD_FOLDER = '../uploads'
ALLOWED_EXTENSIONS = set(['txt', 'pdf', 'docx'])

APP_ROOT = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__)
app.secret_key = "bubblysnowmeninjeanschicago"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def getRandomNumber(n):
    global files

    while True:
        num = int(''.join(random.sample('0123456789', n)))

        if num not in files:
            files.append(num)
            return num
    

@app.route('/')
def index():
    return render_template("index.html")


@app.route('/', methods=['POST'])
def get_text():
    summ = []
    if request.method == 'POST':

        if 'text' in request.form:
            text = request.form.get('text')
            size = int(request.form.get('copy-size'))

            summ = summarize(text, n=size)

        elif 'url' in request.form:
            url = request.form.get('url')
            print("MY URL: " + url)
            # title, text = getArticle(url)
            title = session['url-title']
            text = session['url-text']
            size = int(request.form.get('url-size'))

            summ = summarize(text, title, n=size)

        elif 'placeholder' in request.form:
            target = os.path.join(APP_ROOT, "uploads")

            if not os.path.isdir(target):
                os.mkdir(target)

            file = request.files['file']
            size = int(request.form.get('upload-size'))
        
            if file.filename == '':
                flash('No selected file')
                return redirect(request.url)

            if file and allowed_file(file.filename):
                filename = str(getRandomNumber(5)) + "_" + secure_filename(file.filename)
                print("Filename: " + str(filename))
                destination = "\\".join([target, filename])
                print(destination)
                file.save(destination)
                session['uploaded'] = destination
                text = getText(destination)
                summ = summarize(text, n=size)

            else:
                flash("Unsupported file. Please upload only .docx, .pdf or .txt files.")
                return redirect(request.url)

        session['summary'] = summ
        percentage = getPercentage(text, summ)
        topWords = getTopWords(text)
        return render_template("result.html", summaries=summ, percentage=percentage, words=topWords)

    print("FAILURE")
    return render_template("index.html")


@app.route('/download/')
def download():
    try:
        summ = session['summary']
        target = os.path.join(APP_ROOT, "downloads")
        
        if not os.path.isdir(target):
            os.mkdir(target)
        
        filename = "Summary" + str(getRandomNumber(7)) + ".docx"
        file = Document()
        destination = "\\".join([target, filename])
        print("Download destination: " + destination)

        if session['url-title'] != "":
            file.add_heading(session['url-title'], level=2)
        else:
            file.add_heading("Summary", level=1)

        for sent in summ:
            file.add_paragraph(sent)

        file.save(destination)
        return send_file(destination, attachment_filename=filename, as_attachment=True)
        

    except Exception as e:
        return """
        <p>An unexpected error has occured. Click <a href="/">here</a> to go back.</p>
        Error : 
        """ + str(e)


@app.route('/_show_url_content/')
def show_url_content():
    try:
        url = request.args.get('url', 0, type=str)
        #print("URL: " + url)

        title, text = getArticle(url)
        session['url-title'] = title
        session['url-text'] = text
        return jsonify(title=title, result=text)

    except Exception as e:
        return str(e)


if __name__ == "__main__":
    app.run(debug=True)